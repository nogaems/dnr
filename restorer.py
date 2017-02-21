#!/usr/bin/env python3
# coding=utf8

import argparse
import logging as log
import os
import sys
import re
import subprocess
import shutil
# external dependencies
import docx


parser = argparse.ArgumentParser(
    formatter_class=argparse.RawDescriptionHelpFormatter)

parser.add_argument('--input', '-i', action='store',
                    help='Path to the directory that contains document files that you was restored')
parser.add_argument('--output', '-o', action='store',
                    help='Specified path to save.')
parser.add_argument('--verbose', '-v', action='store_true',
                    help='Verbose output')

args = parser.parse_args()

if args.verbose:
    log.basicConfig(format="%(levelname)s: %(message)s", level=log.INFO)
    log.info('Verbose output.')
    convert_output = sys.stdout
else:
    log.basicConfig(format="%(levelname)s: %(message)s")
    convert_output = open(os.devnull)

input_path = os.path.abspath(
    os.path.curdir) if not args.input else os.path.abspath(args.input)
output_path = os.path.abspath(
    os.path.curdir) if not args.output else os.path.abspath(args.output)

log.info('Input path: {}'.format(input_path))
log.info('Output path: {}'.format(output_path))

formats = ['doc', 'docx']
files = [f for f in os.listdir(
    input_path) if re.match(r'^.*\.(doc|docx)+$', f)]

log.info('Found {} files'.format(len(files)))
if len(files) > 0:
    if not os.path.isdir(output_path):
        log.info(
            "Output path \'{}\' has not exists yet! Making directory".format(output_path))
        os.mkdir(output_path)
else:
    print("Done!")
    exit()
for old_name in files:
    if old_name.endswith("doc"):
        subprocess.call(
            ['soffice',
             '--headless',
             '--convert-to',
             'docx',
             '--outdir',
             input_path,
             os.path.join(input_path, old_name)],
            stdout=convert_output)
        os.remove(os.path.join(input_path, old_name))
        old_name = old_name + "x"
    document = docx.Document(os.path.join(input_path, old_name))
    if len(document.paragraphs) is 0:
        log.info('File {} has not any paragraphs')
        continue
    for p in document.paragraphs:
        if len(p.text) is 0:
            continue
        base_name = p.text[:64] if len(p.text) >= 64 else p.text[:len(p.text)]
        break

    counter = 0
    new_name = base_name + ".docx"
    while os.path.isfile(os.path.join(output_path, new_name)):
        counter += 1
        new_name = "{}_{}.docx".format(base_name, counter)

    log.info('Moving: {} -> {}'.format(
        os.path.join(input_path, old_name),
        os.path.join(output_path, new_name)
    ))

    shutil.move(os.path.join(input_path, old_name),
                os.path.join(output_path, new_name))

print("Done!")
